# -*- coding: utf-8 -*-
"""
渲染器：把抽取结果渲染成要素式起诉状。

三种输出：
  - render_text()     纯文本（可直接贴入法院系统/打印）
  - render_markdown() Markdown（含表格，便于预览/版本管理）
  - render_docx()     Word .docx（符合法院提交格式：宋体小四、A4、表格外框封闭）

要素式文书表格规范：每张表格【最外层周边封闭】（外圈连续粗边框），内部分隔为细线。
"""

import io
from typing import Dict, List, Tuple

from .schemas import get_schema, COMMON_PARTY_GROUP_PLAINTIFF, COMMON_PARTY_GROUP_DEFENDANT


# ─────────────────────────────────────────────
# 统一数据收集：把 result 转成有序的渲染单元
# ─────────────────────────────────────────────
def _gather(result: Dict, cause: str) -> List[Tuple[str, object, str]]:
    """
    返回 [(group_name, payload, kind), ...]
    kind:
      'table'    payload = [(label, value, required), ...]
      'list'     payload = [str, ...]   （诉讼请求；诉讼费用单独走 result['cost_burden']）
      'evidence' payload = [{name,form,purpose,source}, ...]
    """
    items: List[Tuple[str, object, str]] = []

    def _party_rows(role: str):
        group = COMMON_PARTY_GROUP_PLAINTIFF if role == 'plaintiff' else COMMON_PARTY_GROUP_DEFENDANT
        return [(f.label, result[role].get(f.key, ''), f.required) for f in group.fields]

    items.append(('原告信息', _party_rows('plaintiff'), 'table'))
    items.append(('被告信息', _party_rows('defendant'), 'table'))
    items.append(('诉讼请求', result.get('claims', []) or [], 'list'))

    schema = get_schema(cause)
    if schema:
        for g in schema.fact_groups:
            rows = []
            for f in g.fields:
                label = f"{f.label}（{f.unit}）" if f.unit else f.label
                rows.append((label, result['facts'].get(f.key, ''), f.required))
            items.append((g.name, rows, 'table'))

    items.append(('证据清单', result.get('evidence', []) or [], 'evidence'))
    return items


def _cost_line(result: Dict) -> str:
    """诉讼费用条款：优先用从原文剥离的，否则补默认。保证以句号结尾。"""
    cost = (result.get('cost_burden') or '').strip().strip('。.；;')
    return cost + '。' if cost else '本案诉讼费用由被告承担。'


# ─────────────────────────────────────────────
# 纯文本
# ─────────────────────────────────────────────
def render_text(result: Dict, cause: str) -> str:
    lines = []
    lines.append("民事起诉状".center(28))
    lines.append("（要素式）".center(26))
    lines.append("")
    lines.append(f"案    由：{cause}")
    lines.append("")

    for name, payload, kind in _gather(result, cause):
        lines.append(f"【{name}】")
        if kind == 'table':
            for label, value, required in payload:
                if value:
                    lines.append(f"    {label}：{value}")
                elif required:
                    lines.append(f"    {label}：________________（待填）")
            lines.append("")
        elif kind == 'list':
            if not payload:
                lines.append("    1. ____________________________（待填）")
            else:
                for i, c in enumerate(payload, 1):
                    lines.append(f"    {i}. {c}")
            lines.append(f"    {_cost_line(result)}")
            lines.append("")
        elif kind == 'evidence':
            if not payload:
                lines.append("    1. ____________________________（请补充证据）")
            else:
                for i, e in enumerate(payload, 1):
                    parts = [e.get('name', '')]
                    if e.get('form'):
                        parts.append(f"证据形式：{e['form']}")
                    if e.get('purpose'):
                        parts.append(f"证明对象：{e['purpose']}")
                    if e.get('source'):
                        parts.append(f"来源：{e['source']}")
                    lines.append(f"    {i}. {'；'.join(parts)}")
            lines.append("")

    lines.append("此    致")
    lines.append("          ____________________人民法院")
    lines.append("")
    lines.append("                              具状人（签名/盖章）：")
    lines.append("                                  ______年______月______日")
    return "\n".join(lines)


# ─────────────────────────────────────────────
# Markdown
# ─────────────────────────────────────────────
def render_markdown(result: Dict, cause: str) -> str:
    md = []
    md.append("# 民事起诉状（要素式）\n")
    md.append(f"**案由：** {cause}\n")

    for name, payload, kind in _gather(result, cause):
        md.append(f"\n## {name}\n")
        if kind == 'table':
            md.append("| 字段 | 内容 |")
            md.append("|---|---|")
            for label, value, required in payload:
                if value:
                    md.append(f"| {label} | {value} |")
                elif required:
                    md.append(f"| {label} | _待填_ |")
        elif kind == 'list':
            if not payload:
                md.append("1. _待填_")
            else:
                for i, c in enumerate(payload, 1):
                    md.append(f"{i}. {c}")
            md.append(f"\n{_cost_line(result)}")
        elif kind == 'evidence':
            md.append("| 序号 | 证据名称 | 证据形式 | 证明对象 | 来源 |")
            md.append("|---|---|---|---|---|")
            if not payload:
                md.append("| 1 | _请补充_ | | | |")
            else:
                for i, e in enumerate(payload, 1):
                    md.append(f"| {i} | {e.get('name', '')} | {e.get('form', '')} | {e.get('purpose', '')} | {e.get('source', '')} |")

    md.append("\n---\n")
    md.append("此致  __________人民法院\n")
    md.append("具状人（签名/盖章）：__________   日期：____年__月__日")
    return "\n".join(md)


# ─────────────────────────────────────────────
# Word .docx
# ─────────────────────────────────────────────
def _set_cn_font(run, font_name: str, size_pt: float):
    """设置中文字体（python-docx 需显式设置 eastAsia）。"""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def _apply_table_borders(table, outer_sz: int = 12, inner_sz: int = 4):
    """设置表格边框：外圈周边粗线（封闭），内部细线。

    outer_sz/inner_sz 单位为 1/8 pt（Word 边框尺寸单位）。
    outer_sz=12 → 1.5pt 粗外框；inner_sz=4 → 0.5pt 细内线。
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tbl = table._tbl
    tblPr = tbl.tblPr
    # 移除已有 tblBorders，避免叠加
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    borders = OxmlElement('w:tblBorders')
    # 必须按 W3C 顺序：top,left,bottom,right,insideH,insideV
    for edge, sz in (('top', outer_sz), ('left', outer_sz),
                     ('bottom', outer_sz), ('right', outer_sz),
                     ('insideH', inner_sz), ('insideV', inner_sz)):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(sz))
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000')
        borders.append(b)
    tblPr.append(borders)


def _fill_table_cells_font(table, font_name: str, size_pt: float):
    """给表格所有单元格的文字统一设置中文字体。"""
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    _set_cn_font(run, font_name, size_pt)


def render_docx(result: Dict, cause: str, out_path: str) -> str:
    """渲染为 Word 文档，返回保存路径。"""
    import config
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    # 页面边距收窄，让要素式表格更舒展
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("民事起诉状（要素式）")
    _set_cn_font(r, config.DOCX_TITLE_FONT, config.DOCX_TITLE_SIZE)
    r.bold = True

    # 案由
    p = doc.add_paragraph()
    r = p.add_run(f"案由：{cause}")
    _set_cn_font(r, config.DOCX_FONT, config.DOCX_FONT_SIZE)
    r.bold = True

    for name, payload, kind in _gather(result, cause):
        # 组标题
        h = doc.add_paragraph()
        r = h.add_run(f"【{name}】")
        _set_cn_font(r, config.DOCX_HEADING_FONT, config.DOCX_HEADING_SIZE)
        r.bold = True

        if kind == 'table':
            rows = [(lbl, val, req) for lbl, val, req in payload if val or req]
            tbl = doc.add_table(rows=1 + len(rows), cols=2)
            tbl.style = 'Table Grid'
            _apply_table_borders(tbl)   # 外圈封闭
            # 列宽：标签窄、内容宽
            try:
                tbl.columns[0].width = Cm(4.5)
                tbl.columns[1].width = Cm(11.5)
            except Exception:
                pass
            hdr = tbl.rows[0].cells
            for cell, text in zip(hdr, ['要素', '内容']):
                run = cell.paragraphs[0].add_run(text)
                run.bold = True
            for i, (lbl, val, req) in enumerate(rows, 1):
                cells = tbl.rows[i].cells
                cells[0].text = lbl
                cells[1].text = val if val else ('（待填）' if req else '')
            _fill_table_cells_font(tbl, config.DOCX_FONT, config.DOCX_FONT_SIZE)
        elif kind == 'list':
            if not payload:
                p = doc.add_paragraph('1. ____________________（待填）')
            else:
                for i, c in enumerate(payload, 1):
                    doc.add_paragraph(f"{i}. {c}")
            doc.add_paragraph(_cost_line(result))
        elif kind == 'evidence':
            evids = payload if payload else [{'name': '', 'form': '', 'purpose': '', 'source': ''}]
            tbl = doc.add_table(rows=1 + len(evids), cols=5)   # 序号|名称|形式|证明对象|来源
            tbl.style = 'Table Grid'
            _apply_table_borders(tbl)
            for cell, text in zip(tbl.rows[0].cells, ['序号', '证据名称', '证据形式', '证明对象', '来源']):
                cell.paragraphs[0].add_run(text).bold = True
            for i, e in enumerate(evids, 1):
                cells = tbl.rows[i].cells
                cells[0].text = str(i)
                cells[1].text = e.get('name', '')
                cells[2].text = e.get('form', '')
                cells[3].text = e.get('purpose', '')
                cells[4].text = e.get('source', '')
            _fill_table_cells_font(tbl, config.DOCX_FONT, config.DOCX_FONT_SIZE)

        doc.add_paragraph('')   # 组间空行

    # 结尾
    doc.add_paragraph('此致')
    doc.add_paragraph('          ____________________人民法院')
    doc.add_paragraph('')
    p = doc.add_paragraph('                              具状人（签名/盖章）：')
    if p.runs:
        _set_cn_font(p.runs[0], config.DOCX_FONT, config.DOCX_FONT_SIZE)
    doc.add_paragraph('                                  ______年______月______日')

    doc.save(out_path)
    return out_path


def docx_bytes_of(result: Dict, cause: str) -> bytes:
    """渲染到内存 bytes（供 Web 下载，直接写 BytesIO，不落临时文件）。

    python-docx 的 doc.save 接受 file-like 对象，故 render_docx 的 out_path
    既可以是路径字符串，也可以是 BytesIO。
    """
    buf = io.BytesIO()
    render_docx(result, cause, buf)
    return buf.getvalue()


# ─────────────────────────────────────────────
# 强制执行申请书（要素式）渲染
# result = {'doc_type','computation','applicant','respondent'}
# ─────────────────────────────────────────────
_EXEC_PARTY_FIELDS = [
    ('name', '姓名/名称'),
    ('id', '身份证号/统一社会信用代码'),
    ('addr', '住所'),
    ('phone', '联系电话'),
]


def _exec_party_rows(party: Dict) -> List[Tuple[str, str]]:
    return [(label, party.get(k, '')) for k, label in _EXEC_PARTY_FIELDS]


def _exec_amount_rows(s: Dict):
    """金额明细行：(label, value, is_total)。0 值项由渲染层决定是否隐藏。"""
    return [
        (s.get('principal_label', '本金'), s.get('principal', '0'), False),
        ('利息', s.get('confirmed_interest', '0'), False),
        ('逾期利息', s.get('overdue_interest', '0'), False),
        ('案件受理费', s.get('litigation_fee', '0'), False),
        ('财产保全费', s.get('preservation_fee', '0'), False),
        ('公告费', s.get('announcement_fee', '0'), False),
        ('律师代理费', s.get('attorney_fee', '0'), False),
        ('财产保全担保费', s.get('guarantee_fee', '0'), False),
        ('已付款（已抵扣）', s.get('paid_amount', '0'), False),
        ('合计', s.get('total', '0'), True),
    ]


def _exec_interest_rows(s: Dict):
    return [(lbl, s.get(key, '')) for lbl, key in (
        ('计息基数', 'interest_base'),
        ('利率/计息方式', 'interest_rate_description'),
        ('起算日', 'interest_start_date'),
        ('截止日', 'cutoff_date'),
        ('计息年天数', 'year_days'),
        ('起止日计入方式', 'date_inclusion'),
    )]


def _exec_other_claims(s: Dict) -> List[str]:
    items = []
    if s.get('has_double_interest_clause'):
        items.append('被申请人加倍支付迟延履行期间的债务利息')
    if s.get('joint_liability_text'):
        items.append(s['joint_liability_text'])
    if s.get('supplementary_liability_text'):
        items.append(s['supplementary_liability_text'])
    for t in s.get('priority_execution_clauses') or []:
        items.append(t)
    for t in s.get('manual_review_clauses') or []:
        items.append('【人工核对】' + t)
    return items


def _exec_basis(s: Dict) -> str:
    return f"{s.get('case_number', '')}{s.get('document_name', '')}".strip()


def render_execution_markdown(result: Dict) -> str:
    comp = result['computation']
    s = comp.structured_params
    app = result.get('applicant', {}) or {}
    resp = result.get('respondent', {}) or {}
    md = ["# 强制执行申请书（要素式）\n", f"**执行依据：** {_exec_basis(s)}\n"]

    for title, party in (('申请人信息', app), ('被申请人信息', resp)):
        md.append(f"\n## {title}\n")
        md.append("| 字段 | 内容 |")
        md.append("|---|---|")
        for label, val in _exec_party_rows(party):
            if val:
                md.append(f"| {label} | {val} |")

    md.append('\n## 申请执行事项\n')
    for line in comp.preview_text.split('\n'):
        if line.strip():
            md.append(line.strip())

    md.append('\n## 金额明细\n')
    md.append("| 项目 | 金额（元） |")
    md.append("|---|---|")
    for label, val, is_total in _exec_amount_rows(s):
        if is_total:
            md.append(f"| **{label}** | **{val}** |")
        else:
            md.append(f"| {label} | {val} |")

    interest_rows = [(lbl, v) for lbl, v in _exec_interest_rows(s) if v]
    if interest_rows:
        md.append('\n## 利息计算要素\n')
        md.append("| 字段 | 内容 |")
        md.append("|---|---|")
        for label, val in interest_rows:
            md.append(f"| {label} | {val} |")

    others = _exec_other_claims(s)
    if others:
        md.append('\n## 其他请求\n')
        for i, t in enumerate(others, 1):
            md.append(f"{i}. {t}")

    if comp.warnings:
        md.append('\n## ⚠️ 提示（需人工核对）\n')
        for w in comp.warnings:
            md.append(f"- {w}")

    md.append('\n---\n')
    md.append('此致  __________人民法院\n')
    md.append('申请人（签名/盖章）：__________   日期：____年__月__日')
    return '\n'.join(md)


def render_execution_text(result: Dict) -> str:
    comp = result['computation']
    s = comp.structured_params
    app = result.get('applicant', {}) or {}
    resp = result.get('respondent', {}) or {}
    lines = ['强制执行申请书'.center(28), '（要素式）'.center(26), '',
             f"执行依据：{_exec_basis(s)}", '']

    for title, party in (('申请人信息', app), ('被申请人信息', resp)):
        lines.append(f'【{title}】')
        for label, val in _exec_party_rows(party):
            if val:
                lines.append(f'    {label}：{val}')
        lines.append('')

    lines.append('【申请执行事项】')
    for line in comp.preview_text.split('\n'):
        if line.strip():
            lines.append(f'    {line.strip()}')
    lines.append('')

    lines.append('【金额明细】')
    for label, val, is_total in _exec_amount_rows(s):
        mark = '  ★合计' if is_total else ''
        lines.append(f'    {label}：{val}元{mark}')
    lines.append('')

    interest_rows = [(lbl, v) for lbl, v in _exec_interest_rows(s) if v]
    if interest_rows:
        lines.append('【利息计算要素】')
        for label, val in interest_rows:
            lines.append(f'    {label}：{val}')
        lines.append('')

    others = _exec_other_claims(s)
    if others:
        lines.append('【其他请求】')
        for i, t in enumerate(others, 1):
            lines.append(f'    {i}. {t}')
        lines.append('')

    if comp.warnings:
        lines.append('【提示（需人工核对）】')
        for w in comp.warnings:
            lines.append(f'    ! {w}')
        lines.append('')

    lines += ['此致', '          ____________________人民法院', '',
              '                              申请人（签名/盖章）：',
              '                                  ______年______月______日']
    return '\n'.join(lines)


def render_execution_docx(result: Dict, out_path: str) -> str:
    import config
    from docx import Document
    from docx.shared import Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    comp = result['computation']
    s = comp.structured_params
    app = result.get('applicant', {}) or {}
    resp = result.get('respondent', {}) or {}

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run('强制执行申请书（要素式）')
    _set_cn_font(r, config.DOCX_TITLE_FONT, config.DOCX_TITLE_SIZE)
    r.bold = True

    p = doc.add_paragraph()
    r = p.add_run(f"执行依据：{_exec_basis(s)}")
    _set_cn_font(r, config.DOCX_FONT, config.DOCX_FONT_SIZE)
    r.bold = True

    def _heading(text):
        h = doc.add_paragraph()
        rr = h.add_run(f'【{text}】')
        _set_cn_font(rr, config.DOCX_HEADING_FONT, config.DOCX_HEADING_SIZE)
        rr.bold = True

    def _kv_table(pairs):
        pairs = [(lbl, val) for lbl, val in pairs if val]
        if not pairs:
            return
        tbl = doc.add_table(rows=1 + len(pairs), cols=2)
        tbl.style = 'Table Grid'
        _apply_table_borders(tbl)
        try:
            tbl.columns[0].width = Cm(4.5)
            tbl.columns[1].width = Cm(11.5)
        except Exception:
            pass
        for cell, text in zip(tbl.rows[0].cells, ['要素', '内容']):
            cell.paragraphs[0].add_run(text).bold = True
        for i, (lbl, val) in enumerate(pairs, 1):
            tbl.rows[i].cells[0].text = lbl
            tbl.rows[i].cells[1].text = str(val)
        _fill_table_cells_font(tbl, config.DOCX_FONT, config.DOCX_FONT_SIZE)
        doc.add_paragraph('')

    for ttl, party in (('申请人信息', app), ('被申请人信息', resp)):
        _heading(ttl)
        _kv_table(_exec_party_rows(party))

    _heading('申请执行事项')
    for line in comp.preview_text.split('\n'):
        if line.strip():
            doc.add_paragraph(line.strip())
    doc.add_paragraph('')

    _heading('金额明细')
    amts = _exec_amount_rows(s)
    tbl = doc.add_table(rows=1 + len(amts), cols=2)
    tbl.style = 'Table Grid'
    _apply_table_borders(tbl)
    for cell, text in zip(tbl.rows[0].cells, ['项目', '金额（元）']):
        cell.paragraphs[0].add_run(text).bold = True
    for i, (lbl, val, is_total) in enumerate(amts, 1):
        c0 = tbl.rows[i].cells[0]
        c1 = tbl.rows[i].cells[1]
        c0.text = lbl
        c1.text = str(val)
        if is_total:
            for c in (c0, c1):
                for run in c.paragraphs[0].runs:
                    run.bold = True
    _fill_table_cells_font(tbl, config.DOCX_FONT, config.DOCX_FONT_SIZE)
    doc.add_paragraph('')

    interest_rows = [(lbl, v) for lbl, v in _exec_interest_rows(s) if v]
    if interest_rows:
        _heading('利息计算要素')
        _kv_table(interest_rows)

    others = _exec_other_claims(s)
    if others:
        _heading('其他请求')
        for i, t in enumerate(others, 1):
            doc.add_paragraph(f'{i}. {t}')
        doc.add_paragraph('')

    if comp.warnings:
        _heading('提示（需人工核对）')
        for w in comp.warnings:
            doc.add_paragraph(f'! {w}')
        doc.add_paragraph('')

    doc.add_paragraph('此致')
    doc.add_paragraph('          ____________________人民法院')
    doc.add_paragraph('')
    p = doc.add_paragraph('                              申请人（签名/盖章）：')
    if p.runs:
        _set_cn_font(p.runs[0], config.DOCX_FONT, config.DOCX_FONT_SIZE)
    doc.add_paragraph('                                  ______年______月______日')

    doc.save(out_path)
    return out_path
